import os
import json
import logging
from dotenv import load_dotenv
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
from openai import AzureOpenAI
from config1 import Config
import re
from PIL import Image
import pytesseract
import io
from search_utilities1 import create_search_index, upload_documents, index_on_prem_faiss

# Setup logging
logging.basicConfig(level=logging.INFO)

# Load environment variables from .env file
load_dotenv()

# Get the data folder path and chunk settings from the .env file
data_folder_path = os.getenv('DATA_FOLDER_PATH')
chunk_size = int(os.getenv('CHUNK_SIZE', 1000))  # Default chunk size is 1000 characters
chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))  # Default overlap is 200 characters

config = Config()

# Initialize Azure OpenAI client if in cloud approach
if config.APPROACH == 'cloud':
    client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
                         api_key=os.getenv('AZURE_OPENAI_API_KEY'),
                         api_version='2024-02-15-preview')
else:
    # For on-premises, initialize the sentence transformer model
    model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
    model = SentenceTransformer(model_name)

# Function to read PDF files
def read_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]

                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    logging.warning(f"No text found on page {page_num} of {file_path}. Trying OCR.")
                    
                    # If no text, attempt OCR for images
                    images = page.images
                    for image in images:
                        try:
                            pil_image = Image.open(io.BytesIO(image['data']))
                            ocr_text = pytesseract.image_to_string(pil_image)
                            text += ocr_text
                        except Exception as e:
                            logging.error(f"Error during OCR on page {page_num} of {file_path}: {e}")
            return text
    except Exception as e:
        logging.error(f"Error reading PDF file {file_path}: {e}")
        return ""

# Function to read DOCX files and perform OCR on images
def read_docx_with_images(file_path):
    try:
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

        # Process images in DOCX
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data))
                try:
                    ocr_text = pytesseract.image_to_string(img)
                    text.append(ocr_text)
                except Exception as e:
                    logging.error(f"Error during OCR on image in DOCX file {file_path}: {e}")
        
        if not text:
            logging.warning(f"No text found in DOCX file {file_path}.")
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

# Function to read TXT files
def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read().strip()
            if not text:
                logging.warning(f"No text found in TXT file {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading TXT file {file_path}: {e}")
        return ""

# Class for recursive text chunking
class RecursiveChunker:
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, docs: List[Document]) -> List[Document]:
        if not docs:
            logging.warning("No documents provided for chunking.")
            return []

        try:
            splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                encoding_name="cl100k_base",
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        except AttributeError as e:
            logging.error(f"Error initializing text splitter: {e}")
            return []

        chunked_docs = splitter.split_documents(docs)
        return chunked_docs

# Function to generate embeddings
def generate_embeddings(text_list, model="embedding"):
    embeddings = []
    if config.APPROACH == 'cloud':
        logging.info('Generating embeddings using cloud model...')
        for text in text_list:
            try:
                embedding = client.embeddings.create(
                    input=[text],
                    model=model
                ).data[0].embedding
                embeddings.append(embedding)
            except Exception as e:
                logging.error(f"Error generating embedding for text: {e}")
    else:
        logging.info('Generating embeddings using local model...')
        try:
            embeddings = model.encode(text_list).tolist()
        except Exception as e:
            logging.error(f"Error generating embeddings using local model: {e}")
    return embeddings

# Function to handle multiple file types, chunking, and embeddings
def process_files(data_folder_path, chunk_size, chunk_overlap):
    create_search_index()

    chunker = RecursiveChunker(chunk_size, chunk_overlap)
    processed_documents = []

    for filename in os.listdir(data_folder_path):
        file_path = os.path.join(data_folder_path, filename)
        logging.info(f"Processing file: {filename}")

        try:
            if filename.endswith('.pdf'):
                text = read_pdf(file_path)
            elif filename.endswith('.docx'):
                text = read_docx_with_images(file_path)   # function to handle DOCX with images
            elif filename.endswith('.txt'):
                text = read_txt(file_path)
            else:
                logging.warning(f"Unsupported file format: {filename}")
                continue

            logging.info(f"File: {filename} - Length of text: {len(text)} characters")

            if not text:
                logging.warning(f"No content to process for file: {filename}")
                continue

            documents = [Document(page_content=text)]
            chunks = chunker.chunk_documents(documents)
            chunk_texts = [chunk.page_content for chunk in chunks]
            embeddings = generate_embeddings(chunk_texts)

            if len(chunk_texts) != len(embeddings):
                logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
                continue

            for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
                if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
                    logging.error(f"Embedding at index {i} is not a list of floats.")
                    continue

                sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', ''))
                document = {
                    "id": f"{sanitized_title}_{i}",
                    "title": filename,
                    "content": chunk_text,
                    "contentVector": embedding
                }

                if config.APPROACH == 'cloud':
                    upload_documents([document])
                    logging.info(f"Uploaded document for {filename} to Azure Search.")
                else:
                    index_on_prem_faiss([embedding], [f"{sanitized_title}_{i}"])
                    logging.info(f"Uploaded embedding for {filename} to FAISS.")

                processed_documents.append(document)

        except Exception as e:
            logging.error(f"Error processing file {filename}: {e}")

    logging.info("Processing completed for all documents.")
    return processed_documents

# Example usage
if __name__ == "__main__":
    processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
    logging.info(f"Processed {len(processed_data)} documents.")
    logging.info("Chunks and embeddings are created.")
