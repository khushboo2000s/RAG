import os
import logging
from dotenv import load_dotenv
from pdf2image import convert_from_path
import pytesseract
import PyPDF2
import docx
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import List
from sentence_transformers import SentenceTransformer
from openai import AzureOpenAI
from config1 import Config
import re
from search_utilities1 import create_search_index, upload_documents

# Setup logging
logging.basicConfig(level=logging.INFO)

# Load environment variables
load_dotenv()

# Specify the tesseract executable path if not in PATH
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Adjust this path if necessary

# Function to read PDF files with OCR for images
def read_pdf_with_ocr(file_path):
    try:
        images = convert_from_path(file_path)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image)
        return text
    except Exception as e:
        logging.error(f"Error reading PDF file with OCR {file_path}: {e}")
        return ""

# Function to read DOCX files
def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

# Function to read TXT files
def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read().strip()
            return text
    except Exception as e:
        logging.error(f"Error reading TXT file {file_path}: {e}")
        return ""

# Class for recursive text chunking
class RecursiveChunker:
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, docs: List[Document]) -> List[Document]:
        if not docs:
            logging.warning("No documents provided for chunking.")
            return []

        splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            encoding_name="cl100k_base",
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )

        # Split the documents into chunks
        chunked_docs = splitter.split_documents(docs)
        return chunked_docs

# Function to generate embeddings
def generate_embeddings(text_list, model="embedding"):
    embeddings = []
    if config.APPROACH == 'cloud':
        logging.info('Generating embeddings using cloud model...')
        for text in text_list:
            try:
                embedding = client.embeddings.create(
                    input=[text],
                    model=model
                ).data[0].embedding
                embeddings.append(embedding)
            except Exception as e:
                logging.error(f"Error generating embedding for text: {e}")
    else:
        logging.info('Generating embeddings using local model...')
        try:
            embeddings = model.encode(text_list).tolist()
        except Exception as e:
            logging.error(f"Error generating embeddings using local model: {e}")
    return embeddings

# Function to handle multiple file types, chunking, and embeddings
def process_files(data_folder_path, chunk_size, chunk_overlap):
    create_search_index()

    chunker = RecursiveChunker(chunk_size, chunk_overlap)
    processed_documents = []

    for filename in os.listdir(data_folder_path):
        file_path = os.path.join(data_folder_path, filename)
        logging.info(f"Processing file: {filename}")

        try:
            if filename.endswith('.pdf'):
                text = read_pdf_with_ocr(file_path)
            elif filename.endswith('.docx'):
                text = read_docx(file_path)
            elif filename.endswith('.txt'):
                text = read_txt(file_path)
            else:
                logging.warning(f"Unsupported file format: {filename}")
                continue

            logging.info(f"File: {filename} - Length of text: {len(text)} characters")

            if not text:
                logging.warning(f"No content to process for file: {filename}")
                continue

            documents = [Document(page_content=text)]
            chunks = chunker.chunk_documents(documents)
            chunk_texts = [chunk.page_content for chunk in chunks]

            embeddings = generate_embeddings(chunk_texts)

            if len(chunk_texts) != len(embeddings):
                logging.error(f"Mismatch between number of chunks and embeddings")
                continue

            for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
                sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', ''))

                document = {
                    "id": f"{sanitized_title}_{i}",  
                    "title": filename,
                    "content": chunk_text,
                    "contentVector": embedding
                }

                if config.APPROACH == 'cloud':
                    upload_documents([document])  
                    logging.info(f"Uploaded document for {filename} to Azure Search.")
                else:
                    # For FAISS (on-prem)
                    index_on_prem_faiss([embedding], [f"{sanitized_title}_{i}"])
                    logging.info(f"Uploaded embedding for {filename} to FAISS.")

                processed_documents.append(document)

        except Exception as e:
            logging.error(f"Error processing file {filename}: {e}")

    logging.info("Processing completed for all documents.")
    return processed_documents

# Main execution
if __name__ == "__main__":
    # Load environment variables from .env file
    config = Config()
    data_folder_path = os.getenv('DATA_FOLDER_PATH')
    chunk_size = int(os.getenv('CHUNK_SIZE', 1000))  # Default chunk size is 1000 characters
    chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))  # Default overlap is 200 characters

    # Initialize Azure OpenAI client if in cloud approach
    if config.APPROACH == 'cloud':
        client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
                             api_key='',
                             api_version='2024-02-15-preview')
    else:
        # For on-premises, initialize the sentence transformer model
        model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
        model = SentenceTransformer(model_name)

    processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
    logging.info(f"Processed {len(processed_data)} documents.")
    logging.info("Chunks and embeddings are created.")
