import os
import json
import logging
from dotenv import load_dotenv
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
from openai import AzureOpenAI
from config import Config
from create_embeddings import generate_embeddings
from azure_search_ai_service.upload_documents import upload_documents
import re

# Setup logging
logging.basicConfig(level=logging.INFO)

# Load environment variables from .env file
load_dotenv()

# Get the data folder path and chunk settings from the .env file
data_folder_path = os.getenv('DATA_FOLDER_PATH')
chunk_size = int(os.getenv('CHUNK_SIZE', 1000))  # Default chunk size is 1000 characters
chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))  # Default overlap is 200 characters


config = Config()
# Initialize Azure OpenAI client if in cloud approach
if config.APPROACH == 'cloud':
    client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
                     api_key='',
                     api_version='')

else:
    # For on-premises, initialize the sentence transformer model
    model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
    model = SentenceTransformer(model_name)


def read_pdf(file_path):
    """
    Function to read PDF files
    """
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    logging.warning(f"No text found on page {page_num} of {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading PDF file {file_path}: {e}")
        return ""


def read_docx(file_path):
    """
    Function to read DOCX files
    """
    try:
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        if not text:
            logging.warning(f"No text found in DOCX file {file_path}.")
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return ""


def read_txt(file_path):
    """
    Function to read TXT files
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read().strip()
            if not text:
                logging.warning(f"No text found in TXT file {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading TXT file {file_path}: {e}")
        return ""
    
    
def extract_text_from_file(file_path, filename):
    """
    Extracts text from files based on their type (PDF, DOCX, TXT).
    """
    if filename.endswith('.pdf'):
        return read_pdf(file_path)
    elif filename.endswith('.docx'):
        return read_docx(file_path)
    elif filename.endswith('.txt'):
        return read_txt(file_path)
    else:
        logging.warning(f"Unsupported file format: {filename}")
        return None


class RecursiveChunker:
    """
    Class for recursive text chunking
    """
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, docs: List[Document]) -> List[Document]:
        if not docs:
            logging.warning("No documents provided for chunking.")
            return []

        # Check if the method exists and is available
        try:
            splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                encoding_name="cl100k_base",
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        except AttributeError as e:
            logging.error(f"Error initializing text splitter: {e}")
            return []

        # Split the documents into chunks
        chunked_docs = splitter.split_documents(docs)
        return chunked_docs


def clean_filename(filename):
    """
    Clean the filename for use as a document ID.
    """
    return re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', ''))


def process_and_upload_documents(filename, chunk_texts, embeddings, azure_search_config, doc_id):
    """
    Process and upload documents to the relevant system (Azure/FAISS).
    """
    cleaned_title = clean_filename(filename)

    for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
        # Validate embedding format
        if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
            logging.error(f"Embedding at index {i} is not a list of floats.")
            continue

        logging.info(f"Embedding for {filename}_{i}: {embedding}")

        # Prepare document for upload
        document = {
            "id": str(doc_id),
            "title": filename,
            "content": chunk_text,
            "contentVector": embedding
        }

        if config.APPROACH == 'cloud':
            print('--upload')
            upload_documents([document], azure_search_config)  # Upload to Azure Search
            logging.info(f"Uploaded document for {filename} to Azure Search.")
        else:
            # index_on_prem_faiss([embedding], [f"{sanitized_title}_{i}"])  # Upload to FAISS
            logging.info(f"Uploaded embedding for {filename} to FAISS.")



def process_files(data_folder_path, chunk_size, chunk_overlap, azure_search_config, azure_openai_config):
    """
    Function to process files from a folder, chunk them, generate embeddings,
    and upload to Azure AI Search
    """
    try:
        chunker = RecursiveChunker(chunk_size, chunk_overlap)
        doc_id=0
        for filename in os.listdir(data_folder_path):
            file_path = os.path.join(data_folder_path, filename)
            logging.info(f"Processing file: {filename}")
            doc_id+=1

            # Identify file type and extract text
            text = extract_text_from_file(file_path, filename)
            if text:
                documents = [Document(page_content=text)]
                chunked_docs = chunker.chunk_documents(documents)
                
                if not chunked_docs:    #change
                    logging.warning(f"No chunks created for file: {filename}")   #change
                    continue   #change
            
                chunk_texts = [chunk.page_content for chunk in chunked_docs]

                embeddings = generate_embeddings(chunk_texts, azure_openai_config)
                print(len(embeddings))
                print(len(chunk_texts))
                if len(chunk_texts) == len(embeddings):
                    process_and_upload_documents(filename, chunk_texts, embeddings, azure_search_config, str(doc_id))
                else:
                    logging.error(f"Mismatch between number of chunks and embeddings for {filename}")
                    error_msg = f"Mismatch between number of chunks and embeddings for {filename}"
                    logging.exception(error_msg)
                    return error_msg
            else:
                warning_msg = f"No content to process in file: {filename}"
                logging.warning(warning_msg)
                return warning_msg
        
        success_msg = "All files processed successfully."
        logging.info(success_msg)
        return success_msg
    
    except Exception as e:
        logging.exception(f"Error processing file {filename}: {e}")

