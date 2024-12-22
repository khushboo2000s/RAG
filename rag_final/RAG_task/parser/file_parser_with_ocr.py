# import os
# import logging
# from pypdf import PdfReader
# from docx import Document
# import pytesseract
# from pytesseract import Output
# from pdf2image import convert_from_path
# from langchain.text_splitter import RecursiveCharacterTextSplitter

# # Set up logging
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# # Set up Tesseract path (adjust according to your system)
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Update the path as needed

# # Function to extract text from PDF using pypdf and OCR for images
# def extract_text_from_pdf(file_path):
#     try:
#         reader = PdfReader(file_path)
#         text = ""
#         for page_num, page in enumerate(reader.pages):
#             text += page.extract_text() if page.extract_text() else ""

#             # If no text is found, use OCR on the images
#             if not text.strip():
#                 logging.info(f"No text found on page {page_num + 1}, checking for images.")
#                 images = convert_from_path(file_path)
#                 for image in images:
#                     text += pytesseract.image_to_string(image, output_type=Output.STRING)  # OCR for image-based PDFs
#         return text
#     except Exception as e:
#         logging.error(f"Error extracting text from PDF: {e}")
#         return ""

# # Function to extract text from DOCX using python-docx and OCR for images
# def extract_text_from_docx(file_path):
#     try:
#         doc = Document(file_path)
#         text = "\n".join([para.text for para in doc.paragraphs])
        
#         # Handle images in DOCX if required (not all DOCX files may have images accessible this way)
#         # Add custom image processing if necessary

#         return text
#     except Exception as e:
#         logging.error(f"Error extracting text from DOCX: {e}")
#         return ""

# # Function to extract text from TXT files
# def extract_text_from_txt(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             return file.read()
#     except Exception as e:
#         logging.error(f"Error reading TXT file: {e}")
#         return ""

# # Determine file type and extract text
# def parse_file(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == ".pdf":
#         return extract_text_from_pdf(file_path)
#     elif ext == ".docx":
#         return extract_text_from_docx(file_path)
#     elif ext == ".txt":
#         return extract_text_from_txt(file_path)
#     else:
#         logging.warning(f"Unsupported file format: {ext}")
#         return ""

# # Chunk the extracted text using Recursive Text Splitter
# def chunk_text(text):
#     splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#     return splitter.split_text(text)

# # Process files from the data folder
# def process_files(data_folder):
#     for filename in os.listdir(data_folder):
#         file_path = os.path.join(data_folder, filename)
#         logging.info(f"Processing file: {file_path}")
#         file_content = parse_file(file_path)
#         chunks = chunk_text(file_content)
#         logging.info(f"File: {filename}, Chunks: {len(chunks)}")
#         # You can save or process chunks further here

# if __name__ == "__main__":
#     data_folder_path = "D:/Genai_project/Retrieval Augmented Generation/rag_final/RAG_task/data_files"  # Update your path here
#     process_files(data_folder_path)






# import os
# import logging
# import re
# from dotenv import load_dotenv
# from typing import List
# from langchain.schema import Document
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from pypdf import PdfReader
# from docx import Document as DocxDocument
# import pytesseract
# from pytesseract import Output
# from pdf2image import convert_from_path
# from sentence_transformers import SentenceTransformer
# from openai import AzureOpenAI
# from config import Config
# from search_utilities import create_search_index, upload_documents

# # Set up logging
# logging.basicConfig(level=logging.INFO)

# # Load environment variables from .env file
# load_dotenv()

# # Set up Tesseract path (adjust according to your system)
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Update the path as needed

# # Get the data folder path and chunk settings from the .env file
# data_folder_path = os.getenv('DATA_FOLDER_PATH')
# chunk_size = int(os.getenv('CHUNK_SIZE', 1000))  # Default chunk size is 1000 characters
# chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))  # Default overlap is 200 characters

# config = Config()

# # Initialize Azure OpenAI client if in cloud approach
# if config.APPROACH == 'cloud':
#     client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
#                      api_key="",
#                      api_version='2024-02-15-preview')

# else:
#     # For on-premises, initialize the sentence transformer model
#     model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
#     model = SentenceTransformer(model_name)

# # Function to extract text from PDF using pypdf and OCR for images
# def extract_text_from_pdf(file_path):
#     try:
#         reader = PdfReader(file_path)
#         text = ""
#         for page_num, page in enumerate(reader.pages):
#             text += page.extract_text() if page.extract_text() else ""

#             # If no text is found, use OCR on the images
#             if not text.strip():
#                 logging.info(f"No text found on page {page_num + 1}, checking for images.")
#                 images = convert_from_path(file_path)
#                 for image in images:
#                     text += pytesseract.image_to_string(image, output_type=Output.STRING)  # OCR for image-based PDFs
#         return text
#     except Exception as e:
#         logging.error(f"Error extracting text from PDF: {e}")
#         return ""

# # Function to extract text from DOCX using python-docx and OCR for images
# def extract_text_from_docx(file_path):
#     try:
#         doc = DocxDocument(file_path)
#         text = "\n".join([para.text for para in doc.paragraphs])
#         return text
#     except Exception as e:
#         logging.error(f"Error extracting text from DOCX: {e}")
#         return ""

# # Function to extract text from TXT files
# def extract_text_from_txt(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             return file.read()
#     except Exception as e:
#         logging.error(f"Error reading TXT file: {e}")
#         return ""

# # Class for recursive text chunking
# class RecursiveChunker:
#     def __init__(self, chunk_size: int, chunk_overlap: int):
#         self.chunk_size = chunk_size
#         self.chunk_overlap = chunk_overlap

#     def chunk_documents(self, docs: List[Document]) -> List[Document]:
#         if not docs:
#             logging.warning("No documents provided for chunking.")
#             return []

#         # splitter = RecursiveCharacterTextSplitter(chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
#         # # Split the documents into chunks
#         # chunked_docs = splitter.split_documents(docs)
#         # return chunked_docs



#         # Check if the method exists and is available
#         try:
#             splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
#                 encoding_name="cl100k_base",
#                 chunk_size=self.chunk_size,
#                 chunk_overlap=self.chunk_overlap,
#                 separators=["\n\n", "\n", " ", ""]
#             )
#         except AttributeError as e:
#             logging.error(f"Error initializing text splitter: {e}")
#             return []

#         # Split the documents into chunks
#         chunked_docs = splitter.split_documents(docs)
#         return chunked_docs

# # Function to generate embeddings
# def generate_embeddings(text_list, model="embedding"):
#     embeddings = []
#     if config.APPROACH == 'cloud':
#         logging.info('Generating embeddings using cloud model...')
#         for text in text_list:
#             try:
#                 embedding = client.embeddings.create(
#                     input=[text],
#                     model=model
#                 ).data[0].embedding
#                 embeddings.append(embedding)
#             except Exception as e:
#                 logging.error(f"Error generating embedding for text: {e}")
#     else:
#         logging.info('Generating embeddings using local model...')
#         try:
#             embeddings = model.encode(text_list).tolist()
#         except Exception as e:
#             logging.error(f"Error generating embeddings using local model: {e}")
#     return embeddings

# # Function to handle multiple file types, chunking, and embeddings
# def process_files(data_folder_path, chunk_size, chunk_overlap):
#     # Create the search index at the beginning
#     create_search_index()

#     chunker = RecursiveChunker(chunk_size, chunk_overlap)
#     processed_documents = []

#     for filename in os.listdir(data_folder_path):
#         file_path = os.path.join(data_folder_path, filename)
#         logging.info(f"Processing file: {filename}")

#         try:
#             if filename.endswith('.pdf'):
#                 text = extract_text_from_pdf(file_path)
#             elif filename.endswith('.docx'):
#                 text = extract_text_from_docx(file_path)
#             elif filename.endswith('.txt'):
#                 text = extract_text_from_txt(file_path)
#             else:
#                 logging.warning(f"Unsupported file format: {filename}")
#                 continue

#             # Log the length of text read from the file
#             logging.info(f"File: {filename} - Length of text: {len(text)} characters")

#             if not text:
#                 logging.warning(f"No content to process for file: {filename}")
#                 continue

#             # Convert the text into a Document object
#             documents = [Document(page_content=text)]

#             # Chunk the documents
#             chunks = chunker.chunk_documents(documents)

#             # Extract text from chunks
#             chunk_texts = [chunk.page_content for chunk in chunks]

#             # Generate embeddings for each chunk
#             embeddings = generate_embeddings(chunk_texts)

#             if len(chunk_texts) != len(embeddings):
#                 logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
#                 continue

#             # Collect data for uploading
#             for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
#                 # Ensure embedding is a list of floats
#                 if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
#                     logging.error(f"Embedding at index {i} is not a list of floats.")
#                     continue

#                 logging.info(f"Embedding for {filename}_{i}: {embedding}")

#                 # Sanitize the title to create a valid document ID
#                 sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', ''))

#                 # Prepare document for upload
#                 document = {
#                     "id": f"{sanitized_title}_{i}",  
#                     "title": filename,
#                     "content": chunk_text,
#                     "contentVector": embedding
#                 }

#                 # Create Azure Search Index if using cloud
#                 if config.APPROACH == 'cloud':
#                     upload_documents([document])  # Upload the document directly
#                     logging.info(f"Uploaded document for {filename} to Azure Search.")

#                 else:
#                     # For FAISS (on-prem)
#                     index_on_prem_faiss([embedding], [f"{sanitized_title}_{i}"])  # Assuming indexing function works this way
#                     logging.info(f"Uploaded embedding for {filename} to FAISS.")

#                 # Append to processed documents list
#                 processed_documents.append(document)

#         except Exception as e:
#             logging.error(f"Error processing file {filename}: {e}")

#     logging.info("Processing completed for all documents.")
#     return processed_documents

# # Example usage
# if __name__ == "__main__":
#     processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
#     logging.info(f"Processed {len(processed_data)} documents.")
#     logging.info("Chunks and embeddings are created.")



#####################################################################



import os
import json
import logging
from dotenv import load_dotenv
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
from openai import AzureOpenAI
from config import Config
import re
from search_utilities import create_search_index, upload_documents

from PIL import Image
import pytesseract

# Setup logging
logging.basicConfig(level=logging.INFO)

# Load environment variables from .env file
load_dotenv()

# Set up Tesseract path (adjust according to your system)
#pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Update the path as needed


# Get the data folder path and chunk settings from the .env file
data_folder_path = os.getenv('DATA_FOLDER_PATH')
chunk_size = int(os.getenv('CHUNK_SIZE', 1000))  # Default chunk size is 1000 characters
chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))  # Default overlap is 200 characters

config = Config()

# Initialize Azure OpenAI client if in cloud approach
if config.APPROACH == 'cloud':
    client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
                     api_key='',
                     api_version='2024-02-15-preview')

else:
    # For on-premises, initialize the sentence transformer model
    model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
    model = SentenceTransformer(model_name)

# Function to read PDF files
def read_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    logging.warning(f"No text found on page {page_num} of {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading PDF file {file_path}: {e}")
        return ""

# Function to read DOCX files
def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        if not text:
            logging.warning(f"No text found in DOCX file {file_path}.")
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

# Function to read TXT files
def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read().strip()
            if not text:
                logging.warning(f"No text found in TXT file {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading TXT file {file_path}: {e}")
        return ""



# Function to perform OCR on image files
def read_image(file_path):
    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        if not text:
            logging.warning(f"No text found in image file {file_path}.")
        return text
    except Exception as e:
        logging.error(f"Error reading image file {file_path}: {e}")
        return ""

####################################        

# Class for recursive text chunking
class RecursiveChunker:
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, docs: List[Document]) -> List[Document]:
        if not docs:
            logging.warning("No documents provided for chunking.")
            return []

        # Check if the method exists and is available
        try:
            splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                encoding_name="cl100k_base",
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        except AttributeError as e:
            logging.error(f"Error initializing text splitter: {e}")
            return []

        # Split the documents into chunks
        chunked_docs = splitter.split_documents(docs)
        return chunked_docs

# Function to generate embeddings
def generate_embeddings(text_list, model="embedding"):
    embeddings = []
    if config.APPROACH == 'cloud':
        logging.info('Generating embeddings using cloud model...')
        for text in text_list:
            try:
                embedding = client.embeddings.create(
                    input=[text],
                    model=model  # Use the model parameter as "embedding"
                ).data[0].embedding
                embeddings.append(embedding)
            except Exception as e:
                logging.error(f"Error generating embedding for text: {e}")
    else:
        logging.info('Generating embeddings using local model...')
        try:
            embeddings = model.encode(text_list).tolist()  # Ensure embeddings are list
        except Exception as e:
            logging.error(f"Error generating embeddings using local model: {e}")
    return embeddings

# Function to handle multiple file types, chunking, and embeddings
def process_files(data_folder_path, chunk_size, chunk_overlap):
    #Create the search index at the beginning
    create_search_index()

    chunker = RecursiveChunker(chunk_size, chunk_overlap)
    #file_data = []
    processed_documents = []

    for filename in os.listdir(data_folder_path):
        file_path = os.path.join(data_folder_path, filename)
        logging.info(f"Processing file: {filename}")

        try:
            if filename.endswith('.pdf'):
                text = read_pdf(file_path)
            elif filename.endswith('.docx'):
                text = read_docx(file_path)
            elif filename.endswith('.txt'):
                text = read_txt(file_path)
            else:
                logging.warning(f"Unsupported file format: {filename}")
                continue

            # Log the length of text read from the file
            logging.info(f"File: {filename} - Length of text: {len(text)} characters")

            if not text:
                logging.warning(f"No content to process for file: {filename}")
                continue

            # Convert the text into a Document object
            documents = [Document(page_content=text)]

            # Chunk the documents
            chunks = chunker.chunk_documents(documents)

            # Extract text from chunks
            chunk_texts = [chunk.page_content for chunk in chunks]

            # Generate embeddings for each chunk
            embeddings = generate_embeddings(chunk_texts)

            if len(chunk_texts) != len(embeddings):
                logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
                continue

            # Collect data for uploading
            for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
                # Ensure embedding is a list of floats
                if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
                    logging.error(f"Embedding at index {i} is not a list of floats.")
                    continue

                logging.info(f"Embedding for {filename}_{i}: {embedding}") 

                # Sanitize the title to create a valid document ID
                sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', ''))


#######################################################################################

                # Prepare document for upload
                document = {
                    "id": f"{sanitized_title}_{i}",  
                    "title": filename,
                    "content": chunk_text,
                    "contentVector": embedding
                }

                # Create Azure Search Index if using cloud
                if config.APPROACH == 'cloud':
                    upload_documents([document])  # Upload the document directly
                    logging.info(f"Uploaded document for {filename} to Azure Search.")

                else:
                    # For FAISS (on-prem)
                    index_on_prem_faiss([embedding], [f"{sanitized_title}_{i}"])  # Assuming indexing function works this way
                    logging.info(f"Uploaded embedding for {filename} to FAISS.")

                # Append to processed documents list
                processed_documents.append(document)

        except Exception as e:
            logging.error(f"Error processing file {filename}: {e}")

    logging.info("Processing completed for all documents.")
    return processed_documents

# Example usage
if __name__ == "__main__":
    processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
    logging.info(f"Processed {len(processed_data)} documents.")
    logging.info("Chunks and embeddings are created.")
