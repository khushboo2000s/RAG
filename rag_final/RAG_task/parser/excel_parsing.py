# # import os
# # import logging
# # from dotenv import load_dotenv
# # from typing import List
# # from langchain.schema import Document
# # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # import pandas as pd
# # from sentence_transformers import SentenceTransformer
# # from search_utilities import create_search_index, upload_documents

# import os
# import json
# import logging
# from dotenv import load_dotenv
# from typing import List
# from langchain.schema import Document
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import PyPDF2
# import docx
# from sentence_transformers import SentenceTransformer
# from openai import AzureOpenAI
# from config import Config
# import re
# from search_utilities import create_search_index, upload_documents



# # Setup logging
# logging.basicConfig(level=logging.INFO)

# # Load environment variables from .env file
# load_dotenv()

# data_folder_path = os.getenv('DATA_FOLDER_PATH')
# chunk_size = int(os.getenv('CHUNK_SIZE', 1000))
# chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))

# config = Config()

# # Initialize Azure OpenAI client if in cloud approach
# if config.APPROACH == 'cloud':
#     client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
#                          api_key='',
#                          api_version='2024-02-15-preview')
# else:
#     # For on-premises, initialize the sentence transformer model
#     model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
#     model = SentenceTransformer(model_name)

# # model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
# # model = SentenceTransformer(model_name)

# # Function to read Excel files
# def read_excel(file_path):
#     try:
#         df = pd.read_excel(file_path, engine='openpyxl')
#         text = df.to_string(index=False)
#         if not text:
#             logging.warning(f"No text found in Excel file {file_path}.")
#         return text
#     except Exception as e:
#         logging.error(f"Error reading Excel file {file_path}: {e}")
#         return ""

# # # Class for recursive text chunking
# # class RecursiveChunker:
# #     def __init__(self, chunk_size: int, chunk_overlap: int):
# #         self.chunk_size = chunk_size
# #         self.chunk_overlap = chunk_overlap

# #     def chunk_documents(self, docs: List[Document]) -> List[Document]:
# #         if not docs:
# #             logging.warning("No documents provided for chunking.")
# #             return []

# #         splitter = RecursiveCharacterTextSplitter(
# #             chunk_size=self.chunk_size,
# #             chunk_overlap=self.chunk_overlap,
# #             separators=["\n\n", "\n", " ", ""]
# #         )
# #         return splitter.split_documents(docs)


# # Class for recursive text chunking
# class RecursiveChunker:
#     def __init__(self, chunk_size: int, chunk_overlap: int):
#         self.chunk_size = chunk_size
#         self.chunk_overlap = chunk_overlap

#     def chunk_documents(self, docs: List[Document]) -> List[Document]:
#         if not docs:
#             logging.warning("No documents provided for chunking.")
#             return []

#         # Check if the method exists and is available
#         try:
#             splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
#                 encoding_name="cl100k_base",
#                 chunk_size=self.chunk_size,
#                 chunk_overlap=self.chunk_overlap,
#                 separators=["\n\n", "\n", " ", ""]
#             )
#         except AttributeError as e:
#             logging.error(f"Error initializing text splitter: {e}")
#             return []

#         # Split the documents into chunks
#         chunked_docs = splitter.split_documents(docs)
#         return chunked_docs

# # # Function to generate embeddings
# # def generate_embeddings(text_list):
# #     logging.info('Generating embeddings using local model...')
# #     try:
# #         return model.encode(text_list).tolist()
# #     except Exception as e:
# #         logging.error(f"Error generating embeddings: {e}")
# #         return []


# # Function to generate embeddings
# def generate_embeddings(text_list, model="embedding"):
#     embeddings = []
#     if config.APPROACH == 'cloud':
#         logging.info('Generating embeddings using cloud model...')
#         for text in text_list:
#             try:
#                 embedding = client.embeddings.create(
#                     input=[text],
#                     model=model  # Use the model parameter as "embedding"
#                 ).data[0].embedding
#                 embeddings.append(embedding)
#             except Exception as e:
#                 logging.error(f"Error generating embedding for text: {e}")
#     else:
#         logging.info('Generating embeddings using local model...')
#         try:
#             embeddings = model.encode(text_list).tolist()  # Ensure embeddings are list
#         except Exception as e:
#             logging.error(f"Error generating embeddings using local model: {e}")
#     return embeddings


# # # Function to process Excel files
# # def process_excel_files(data_folder_path, chunk_size, chunk_overlap):
# #     create_search_index()
# #     chunker = RecursiveChunker(chunk_size, chunk_overlap)
# #     processed_documents = []

# #     for filename in os.listdir(data_folder_path):
# #         file_path = os.path.join(data_folder_path, filename)
# #         if filename.endswith(('.xlsx', '.xls')):
# #             logging.info(f"Processing Excel file: {filename}")

# #             try:
# #                 text = read_excel(file_path)
# #                 if not text:
# #                     logging.warning(f"No content to process for file: {filename}")
# #                     continue

# #                 documents = [Document(page_content=text)]
# #                 chunks = chunker.chunk_documents(documents)

# #                 logging.info(f"File: {filename} - Number of chunks created: {len(chunks)}")

# #                 chunk_texts = [chunk.page_content for chunk in chunks]
# #                 embeddings = generate_embeddings(chunk_texts)

# #                 if len(chunk_texts) != len(embeddings):
# #                     logging.error(f"Mismatch between number of chunks and embeddings for file: {filename}")
# #                     continue

# #                 for chunk_text, embedding in zip(chunk_texts, embeddings):
# #                     document = {
# #                         "title": filename,
# #                         "content": chunk_text,
# #                         "contentVector": embedding
# #                     }
# #                     processed_documents.append(document)

# #                 upload_documents(processed_documents)
# #                 logging.info(f"Uploaded document for {filename} to Azure Search.")

# #             except Exception as e:
# #                 logging.error(f"Error processing file {filename}: {e}")

# #     logging.info("Processing completed for all Excel files.")
# #     return processed_documents




# # Function to handle multiple file types, chunking, and embeddings
# def process_files(data_folder_path, chunk_size, chunk_overlap):
#     create_search_index()
#     chunker = RecursiveChunker(chunk_size, chunk_overlap)
#     processed_documents = []

#     for filename in os.listdir(data_folder_path):
#         file_path = os.path.join(data_folder_path, filename)
#         logging.info(f"Processing file: {filename}")

#         try:
#             if filename.endswith('.pdf'):
#                 text = read_pdf(file_path)
#             elif filename.endswith('.docx'):
#                 text = read_docx(file_path)
#             elif filename.endswith('.txt'):
#                 text = read_txt(file_path)
#             elif filename.endswith('.csv'):
#                 text = read_csv(file_path)
#             elif filename.endswith('.xlsx'):
#                 text = read_excel(file_path)
#             else:
#                 logging.warning(f"Unsupported file format: {filename}")
#                 continue

#             logging.info(f"File: {filename} - Length of text: {len(text)} characters")

#             if not text:
#                 logging.warning(f"No content to process for file: {filename}")
#                 continue

#             documents = [Document(page_content=text)]
#             chunks = chunker.chunk_documents(documents)

#             logging.info(f"File: {filename} - Number of chunks created: {len(chunks)}")

#             if not chunks:
#                 logging.warning(f"No chunks created for file: {filename}")
#                 continue


#             # Print the chunks for debugging
#             for i, chunk in enumerate(chunks):
#                 logging.info(f"Chunk {i + 1}: {chunk.page_content[:100]}...")  # Show first 100 characters



#             chunk_texts = [chunk.page_content for chunk in chunks]
#             embeddings = generate_embeddings(chunk_texts)
#             logging.info(f"File: {filename} - Generated embeddings for {len(embeddings)} chunks.")

#             if len(chunk_texts) != len(embeddings):
#                 logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
#                 continue


#             # Print embeddings for the first few chunks
#             for i, embedding in enumerate(embeddings[:5]):  # Show first 5 embeddings for debugging
#                 logging.info(f"Embedding {i + 1}: {embedding[:10]}...")  # Show first 10 dimensions



#             sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '').replace('.csv', '').replace('.xlsx', ''))
#             document_id = sanitized_title

#             for chunk_text, embedding in zip(chunk_texts, embeddings):
#                 if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
#                     logging.error(f"Embedding is not a list of floats for chunk of file: {filename}.")
#                     continue

#                 document = {
#                     "id": document_id,
#                     "title": filename,
#                     "content": chunk_text,
#                     "contentVector": embedding
#                 }
#                 processed_documents.append(document)

#             if config.APPROACH == 'cloud':
#                 upload_documents([document])
#                 logging.info(f"Uploaded document for {filename} to Azure Search.")
#             else:
#                 index_on_prem_faiss([embedding], [document_id])
#                 logging.info(f"Uploaded embedding for {filename} to FAISS.")

#         except Exception as e:
#             logging.error(f"Error processing file {filename}: {e}")

#     logging.info("Processing completed for all documents.")
#     return processed_documents

# # Main execution
# if __name__ == "__main__":
#     processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
#     if processed_data:
#         logging.info("Chunks and embeddings are created.")
#     else:
#         logging.error("No data was processed. Check input files or processing logic.")


# # # Main execution
# # if __name__ == "__main__":
# #     processed_data = process_excel_files(data_folder_path, chunk_size, chunk_overlap)
# #     if processed_data:
# #         logging.info("Chunks and embeddings are created for Excel files.")
# #     else:
# #         logging.error("No Excel files were processed. Check the input directory.")





import os
import json
import numpy as np 
import logging
from dotenv import load_dotenv
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
import docx
import pandas as pd
from sentence_transformers import SentenceTransformer
from openai import AzureOpenAI
from config import Config
import re
# from search_utilities import create_search_index, upload_documents
from search_utilities_onprem import create_search_index, upload_documents

# Setup logging
# logging.basicConfig(level=logging.INFO)

# Load environment variables from .env file
load_dotenv()


# Set up logging configuration
log_file_path = 'application.log'  # Path to log file
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file_path, mode='w'),  # Log file handler
        logging.StreamHandler()  # Console handler
    ]
)

# Create logger
logging = logging.getLogger(__name__)

# Get the data folder path and chunk settings from the .env file
data_folder_path = os.getenv('DATA_FOLDER_PATH')
chunk_size = int(os.getenv('CHUNK_SIZE', 1000))  # Default chunk size is 1000 characters
chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 200))  # Default overlap is 200 characters

config = Config()

# Initialize Azure OpenAI client if in cloud approach
if config.APPROACH == 'cloud':
    client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
                         api_key='',
                         api_version='2024-02-15-preview')
else:
    # For on-premises, initialize the sentence transformer model
    model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
    model = SentenceTransformer(model_name)

# Function to read PDF files
def read_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    logging.warning(f"No text found on page {page_num} of {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading PDF file {file_path}: {e}")
        return ""

# Function to read DOCX files
def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        if not text:
            logging.warning(f"No text found in DOCX file {file_path}.")
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

# Function to read TXT files
def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read().strip()
            if not text:
                logging.warning(f"No text found in TXT file {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading TXT file {file_path}: {e}")
        return ""

# Function to read CSV files
def read_csv(file_path):
    try:
        df = pd.read_csv(file_path)
        text = df.to_string(index=False)
        if not text:
            logging.warning(f"No text found in CSV file {file_path}.")
        return text
    except Exception as e:
        logging.error(f"Error reading CSV file {file_path}: {e}")
        return ""

# Function to read Excel files
# def read_excel(file_path):
#     try:
#         df = pd.read_excel(file_path, engine='openpyxl')
#         text = df.to_string(index=False)
#         if not text:
#             logging.warning(f"No text found in Excel file {file_path}.")
#         return text
#     except Exception as e:
#         logging.error(f"Error reading Excel file {file_path}: {e}")
#         return ""


def read_excel(file_path):
    try:
        excel_data = pd.ExcelFile(file_path, engine='openpyxl')
        all_text = []

        for sheet_name in excel_data.sheet_names:
            logging.info(f"Processing sheet: {sheet_name}")
            sheet_df = excel_data.parse(sheet_name)

            if sheet_df.empty:
                logging.warning(f"Sheet '{sheet_name}' in {file_path} is empty.")
                continue

            # Convert sheet content into a string
            sheet_text = sheet_df.to_string(index=False)
            all_text.append(sheet_text)

        # Combine all sheet texts into one
        combined_text = "\n".join(all_text)
        if not combined_text.strip():
            logging.warning(f"No text found in Excel file {file_path}.")
        return combined_text

    except Exception as e:
        logging.error(f"Error reading Excel file {file_path}: {e}")
        return ""


# Class for recursive text chunking
class RecursiveChunker:
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, docs: List[Document]) -> List[Document]:
        if not docs:
            logging.warning("No documents provided for chunking.")
            return []

        try:
            splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                encoding_name="cl100k_base",
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        except AttributeError as e:
            logging.error(f"Error initializing text splitter: {e}")
            return []

        chunked_docs = splitter.split_documents(docs)
        return chunked_docs

# Function to generate embeddings
def generate_embeddings(text_list):
    # from indexing_testing import index_on_prem_faiss    # Lazy import
    # from indexing_onprem import index_on_prem_faiss      #*
    from index_onprem_test import index_on_prem_faiss
    embeddings = []
    if config.APPROACH == 'cloud':
        logging.info('Generating embeddings using cloud model...')
        for text in text_list:
            try:
                embedding = client.embeddings.create(
                    input=[text],
                    model="embedding"  # Use the default embedding model in the cloud
                ).data[0].embedding
                embeddings.append(embedding)
            except Exception as e:
                logging.error(f"Error generating embedding for text: {e}")
    else:
        logging.info('Generating embeddings using local model...')
        try:
            # Ensure the text_list is a list of strings
            if not all(isinstance(text, str) for text in text_list):
                raise ValueError("Each element in text_list should be a string.")

            # Generate embeddings using the local model
            embeddings = model.encode(text_list)  # Example local model output
            
            print(embeddings)  # Check if the output is a list of float arrays

            # Ensure that the embeddings are in a list of float format
            embeddings = [list(map(float, embedding)) for embedding in embeddings]

            # Log the output to check the embedding format
            logging.info(f"Generated embeddings for {len(text_list)} chunks.")

            # Validate the embedding dimensions
            for i, embedding in enumerate(embeddings):
                if len(embedding) != 768:  # Check if embedding length is 768
                    logging.error(f"Invalid embedding dimensions for chunk {i}. Expected 768, but got {len(embedding)}.")
                    continue  # Skip or handle this chunk differently

        except Exception as e:
            logging.error(f"Error generating embeddings using local model: {e}")


    ############ Testing-embeddings ########
    # Validate embedding dimensions and ensure they're 1536 dimensions
    # for embedding in embeddings:
    #     if len(embedding) != 768:   #1536
    #         logging.error(f"Embedding dimension mismatch. Expected 768, but got {len(embedding)}.")
    

    # return embeddings


    # for i, embedding in enumerate(embeddings):
    #     if not isinstance(embedding, list) or len(embedding) != 768:
    #         logging.error(f"Invalid embedding dimensions for chunk {i}. Expected 768, but got {len(embedding)}.")
    #         continue

    return embeddings




# # Function to handle multiple file types, chunking, and embeddings
# def process_files(data_folder_path, chunk_size, chunk_overlap):
#     create_search_index()
#     chunker = RecursiveChunker(chunk_size, chunk_overlap)
#     processed_documents = []

#     for filename in os.listdir(data_folder_path):
#         file_path = os.path.join(data_folder_path, filename)
#         logging.info(f"Processing file: {filename}")

#         try:
            # if filename.endswith('.pdf'):
            #     text = read_pdf(file_path)
            # elif filename.endswith('.docx'):
            #     text = read_docx(file_path)
            # elif filename.endswith('.txt'):
            #     text = read_txt(file_path)
            # elif filename.endswith('.csv'):
            #     text = read_csv(file_path)
            # elif filename.endswith('.xlsx'):
            #     text = read_excel(file_path)
#             else:
#                 logging.warning(f"Unsupported file format: {filename}")
#                 continue

#             logging.info(f"File: {filename} - Length of text: {len(text)} characters")

#             if not text:
#                 logging.warning(f"No content to process for file: {filename}")
#                 continue

#             documents = [Document(page_content=text)]
#             chunks = chunker.chunk_documents(documents)

#             logging.info(f"File: {filename} - Number of chunks created: {len(chunks)}")

#             if not chunks:
#                 logging.warning(f"No chunks created for file: {filename}")
#                 continue

#             chunk_texts = [chunk.page_content for chunk in chunks]
#             embeddings = generate_embeddings(chunk_texts)
#             logging.info(f"File: {filename} - Generated embeddings for {len(embeddings)} chunks.")

#             if len(chunk_texts) != len(embeddings):
#                 logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
#                 continue

#             sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '').replace('.csv', '').replace('.xlsx', ''))
#             document_id = sanitized_title

#             for chunk_text, embedding in zip(chunk_texts, embeddings):
#                 if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
#                     logging.error(f"Embedding is not a list of floats for chunk of file: {filename}.")
#                     continue

#                 document = {
#                     "id": document_id,
#                     "title": filename,
#                     "content": chunk_text,
#                     "contentVector": embedding
#                 }
#                 processed_documents.append(document)

#             if config.APPROACH == 'cloud':
#                 upload_documents([document])
#                 logging.info(f"Uploaded document for {filename} to Azure Search.")
#             else:
#                 index_on_prem_faiss([embedding], [document_id])
#                 logging.info(f"Uploaded embedding for {filename} to FAISS.")

#         except Exception as e:
#             logging.error(f"Error processing file {filename}: {e}")

#     logging.info("Processing completed for all documents.")
#     return processed_documents



# Function to handle multiple file types, chunking, and embeddings
def process_files(data_folder_path, chunk_size, chunk_overlap):
    # from indexing_testing import index_on_prem_faiss      # Lazy import
    # from indexing_onprem import index_on_prem_faiss          #*
    from index_onprem_test import index_on_prem_faiss


    document_mapping = {}


    #Create the search index at the beginning
    create_search_index()

    chunker = RecursiveChunker(chunk_size, chunk_overlap)
    #file_data = []
    processed_documents = []

    for filename in os.listdir(data_folder_path):
        file_path = os.path.join(data_folder_path, filename)
        logging.info(f"Processing file: {filename}")

        try:
            if filename.endswith('.pdf'):
                text = read_pdf(file_path)
            elif filename.endswith('.docx'):
                text = read_docx(file_path)
            elif filename.endswith('.txt'):
                text = read_txt(file_path)
            elif filename.endswith('.csv'):
                text = read_csv(file_path)
            elif filename.endswith('.xlsx') or filename.endswith('.xls'):
                text = read_excel(file_path)
            else:
                logging.warning(f"Unsupported file format: {filename}")
                continue

            # Log the length of text read from the file
            logging.info(f"File: {filename} - Length of text: {len(text)} characters")

            if not text:
                logging.warning(f"No content to process for file: {filename}")
                continue

            # Convert the text into a Document object
            documents = [Document(page_content=text)]
            print(f"Total rows before chunking: {len(documents)}")

            # Chunk the documents
            chunks = chunker.chunk_documents(documents)
            logging.info(f"File: {filename} - Number of chunks created: {len(chunks)}")

            print(f"Number of chunks created: {len(chunks)}")

            if not chunks:
                logging.warning(f"No chunks created for file: {filename}")
                continue

            # Extract text from chunks
            chunk_texts = [chunk.page_content for chunk in chunks]
            for i, chunk_text in enumerate(chunk_texts):
                print(f"Chunk {i}: {chunk_text[:100]}...")  # First 100 characters 


            # Generate embeddings for each chunk
            embeddings = generate_embeddings(chunk_texts)
            logging.info(f"File: {filename} - Generated embeddings for {len(embeddings)} chunks.")
            # for i, embeddings in enumerate(embeddings):
            #     print(f"--Embedding-- {i}: {embeddings[:100]}...")  # First 100 characters 


            if len(chunk_texts) != len(embeddings):
                logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
                continue

            # Validate embedding dimensions   ### Changes ###
            # for embedding in embeddings:
            #     if len(embedding) != 1536:   #768
            #         logging.error(f"Embedding dimension mismatch: {len(embedding)} for {filename}")

 ################## changes ################################
            # # Collect data for uploading
            # for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
            #     # Ensure embedding is a list of floats
            #     if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
            #         logging.error(f"Embedding at index {i} is not a list of floats.")
            #         continue


            #     # Log the dimension of the embedding
            #     embedding_np = np.array(embedding)
            #     logging.info(f"Embedding dimension for {filename}_{i}: {embedding_np.shape[0]}")  # Log the size

            #     if embedding_np.shape[0] != 1536:  #768
            #         logging.error(f"Embedding for {filename}_{i} has invalid dimension: {embedding_np.shape[0]}. Expected: 1536.")   #768
            #         continue

            #     # logging.info(f"Document ID: {document['id']} - Embedding format: {type(embedding)} with length: {len(embedding)}")  #

            #     # logging.info(f"Embedding for {filename}_{i}: {embedding}") 

            #     # Sanitize the title to create a valid document ID
            #     sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', ''))

##### Changes ####################################################
            # Validate embedding dimensions and process chunks
            # for i, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):
            #     if len(embedding) != 768:  #1536
            #         logging.error(f"Invalid embedding dimensions for {filename}_{i}. Expected 768.")
            #         continue

            if isinstance(embeddings, (list, np.ndarray)):  # Check if embedding is a list or numpy array
                if len(embeddings) != 768:
                    logging.error(f"Invalid embedding dimensions for chunk {i}. Expected 768, but got {len(embeddings)}.")
            else:
                logging.error(f"Embedding is not an iterable (chunk {i}). Got type {type(embeddings)}.")


            

                sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '').replace('.csv', '').replace('.xlsx', '').replace('.xls', ''))


#######################################################################################

                # Prepare document for upload
                # document_id = f"{sanitized_title}_{i}_{filename}"
                document_id = sanitized_title

                for chunk_text, embedding in zip(chunk_texts, embeddings):
                    if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
                        logging.error(f"Embedding is not a list of floats for chunk of file: {filename}.")
                        continue

                document = {
                    # "id": f"{sanitized_title}_{i}",  
                    "id" : document_id,
                    "title": filename,
                    "content": chunk_text,
                    "contentVector": embeddings
                }
                processed_documents.append(document)

                # Add to document_mapping

                document_mapping[document_id] = {"title": filename, "content": chunk_text}
                # if document_id in document_mapping:
                #     logging.warning(f"Duplicate document ID detected: {document_id}. Overwriting may occur.")


                # Log addition to verify
                logging.info(f"Added document mapping for ID: {document_id} from file: {filename}")


                # logging.info(f"Document ID: {document['id']} - Embedding format: {type(embedding)} with length: {len(embedding)}")  #

                logging.info(f"Preparing to upload document: {document}")    #

                logging.info(f"Final document mapping: {document_mapping}")


                # Create Azure Search Index if using cloud
                if config.APPROACH == 'cloud':
                    upload_documents([document])  # Upload the document directly
                    logging.info(f"Uploaded document for {filename} to Azure Search.")

                else:
                    # For FAISS (on-prem)
                    # index_on_prem_faiss([embedding], [f"{sanitized_title}_{i}"])  # 
                    # index_on_prem_faiss([embedding], [document['id']])  #
                    index_on_prem_faiss([embedding], [document['id']], document_mapping)

                    logging.info(f"Uploaded embedding for {filename} to FAISS.")

                # Append to processed documents list
                processed_documents.append(document)

        except Exception as e:
            logging.error(f"Error processing file {filename}: {e}", exc_info=True)

    logging.info("Processing completed for all documents.")

    logging.info(f"Total processed documents before return: {len(processed_documents)}")  #


    return processed_documents



# Main execution
if __name__ == "__main__":
    processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
    if processed_data:
        logging.info("Chunks and embeddings are created.")
    else:
        logging.error("No data was processed. Check input files or processing logic.")
