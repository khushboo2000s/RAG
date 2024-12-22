

# ############################################################################
# import os
# import json
# import logging
# from dotenv import load_dotenv
# from typing import List
# from langchain.schema import Document
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# import PyPDF2
# import docx
# import pandas as pd  # Import pandas for handling Excel files
# from sentence_transformers import SentenceTransformer
# from openai import AzureOpenAI
# from config import Config
# import re
# from search_utilities import create_search_index, upload_documents

# # Setup logging
# logging.basicConfig(level=logging.INFO)

# # Load environment variables from .env file
# load_dotenv()

# # Get the data folder path and chunk settings from the .env file
# data_folder_path = os.getenv('DATA_FOLDER_PATH')
# chunk_size = int(os.getenv('CHUNK_SIZE', 1500))  # Default chunk size is 1000 characters
# chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 300))  # Default overlap is 200 characters

# config = Config()

# # Initialize Azure OpenAI client if in cloud approach
# if config.APPROACH == 'cloud':
#     client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
#                          api_key='',
#                          api_version='2024-02-15-preview')
# else:
#     # For on-premises, initialize the sentence transformer model
#     model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
#     model = SentenceTransformer(model_name)

# # Function to read PDF files
# def read_pdf(file_path):
#     try:
#         with open(file_path, 'rb') as file:
#             pdf_reader = PyPDF2.PdfReader(file)
#             text = ""
#             for page_num in range(len(pdf_reader.pages)):
#                 page = pdf_reader.pages[page_num]
#                 page_text = page.extract_text()
#                 if page_text:
#                     text += page_text
#                 else:
#                     logging.warning(f"No text found on page {page_num} of {file_path}.")
#             return text
#     except Exception as e:
#         logging.error(f"Error reading PDF file {file_path}: {e}")
#         return ""

# # Function to read DOCX files
# def read_docx(file_path):
#     try:
#         doc = docx.Document(file_path)
#         text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
#         if not text:
#             logging.warning(f"No text found in DOCX file {file_path}.")
#         return '\n'.join(text)
#     except Exception as e:
#         logging.error(f"Error reading DOCX file {file_path}: {e}")
#         return ""

# # Function to read TXT files
# def read_txt(file_path):
#     try:
#         with open(file_path, 'r', encoding='utf-8') as file:
#             text = file.read().strip()
#             if not text:
#                 logging.warning(f"No text found in TXT file {file_path}.")
#             return text
#     except Exception as e:
#         logging.error(f"Error reading TXT file {file_path}: {e}")
#         return ""

# # # Function to read Excel files
# # def read_excel(file_path):
# #     try:
# #         # Read the Excel file using pandas
# #         excel_data = pd.ExcelFile(file_path)
# #         text = ""
# #         for sheet_name in excel_data.sheet_names:  # Iterate through all sheets
# #             df = excel_data.parse(sheet_name)
# #             sheet_text = df.to_string(index=False)  # Convert the DataFrame to text
# #             text += sheet_text
# #         if not text:
# #             logging.warning(f"No text found in Excel file {file_path}.")
# #         return text
# #     except Exception as e:
# #         logging.error(f"Error reading Excel file {file_path}: {e}")
# #         return ""

# ##########################################
# def read_excel(file_path):
#     try:
#         excel_data = pd.ExcelFile(file_path)
#         text = ""
#         for sheet_name in excel_data.sheet_names:
#             df = excel_data.parse(sheet_name)
#             sheet_text = "\n".join(df.astype(str).apply(' '.join, axis=1))  # Convert each row to a string
#             text += f"Sheet: {sheet_name}\n{sheet_text}\n"
#         return text
#     except Exception as e:
#         logging.error(f"Error reading Excel file {file_path}: {e}")
#         return ""
    
# ###################################################

# # Class for recursive text chunking
# class RecursiveChunker:
#     def __init__(self, chunk_size: int, chunk_overlap: int):
#         self.chunk_size = chunk_size
#         self.chunk_overlap = chunk_overlap

#     def chunk_documents(self, docs: List[Document]) -> List[Document]:
#         if not docs:
#             logging.warning("No documents provided for chunking.")
#             return []

#         # Check if the method exists and is available
#         try:
#             splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
#                 encoding_name="cl100k_base",
#                 chunk_size=self.chunk_size,
#                 chunk_overlap=self.chunk_overlap,
#                 separators=["\n\n", "\n", " ", ""]
#             )
#         except AttributeError as e:
#             logging.error(f"Error initializing text splitter: {e}")
#             return []

#         # Split the documents into chunks
#         chunked_docs = splitter.split_documents(docs)
#         return chunked_docs

# # Function to generate embeddings
# def generate_embeddings(text_list, model="embedding"):
#     embeddings = []
#     if config.APPROACH == 'cloud':
#         logging.info('Generating embeddings using cloud model...')
#         for text in text_list:
#             try:
#                 embedding = client.embeddings.create(
#                     input=[text],
#                     model=model  # Use the model parameter as "embedding"
#                 ).data[0].embedding
#                 embeddings.append(embedding)
#             except Exception as e:
#                 logging.error(f"Error generating embedding for text: {e}")
#     else:
#         logging.info('Generating embeddings using local model...')
#         try:
#             embeddings = model.encode(text_list).tolist()  # Ensure embeddings are list
#         except Exception as e:
#             logging.error(f"Error generating embeddings using local model: {e}")
#     return embeddings

# ####################################################
# # Function to store chunks and embeddings in a separate file
# def save_chunks_and_embeddings(chunks, embeddings, filename):
#     try:
#         output = []
#         for chunk, embedding in zip(chunks, embeddings):
#             chunk_data = {
#                 'chunk': chunk,
#                 'embedding': embedding
#             }
#             output.append(chunk_data)

#         # Save to a JSON file
#         output_file = os.path.join(data_folder_path, f"{filename}_chunks_embeddings.json")
#         with open(output_file, 'w') as f:
#             json.dump(output, f, indent=4)

#         logging.info(f"Saved chunks and embeddings to {output_file}")
#     except Exception as e:
#         logging.error(f"Error saving chunks and embeddings: {e}")

# ###################################################

# # Function to handle multiple file types, chunking, and embeddings
# def process_files(data_folder_path, chunk_size, chunk_overlap):
#     # Create the search index at the beginning
#     create_search_index()

#     chunker = RecursiveChunker(chunk_size, chunk_overlap)
#     processed_documents = []  # Create a list to hold processed documents

#     for filename in os.listdir(data_folder_path):
#         file_path = os.path.join(data_folder_path, filename)
#         logging.info(f"Processing file: {filename}")

#         try:
#             # Handle different file types
#             if filename.endswith('.pdf'):
#                 text = read_pdf(file_path)
#             elif filename.endswith('.docx'):
#                 text = read_docx(file_path)
#             elif filename.endswith('.txt'):
#                 text = read_txt(file_path)
#             elif filename.endswith('.xlsx'):  # Check if it's an Excel file
#                 text = read_excel(file_path)
#             else:
#                 logging.warning(f"Unsupported file format: {filename}")
#                 continue

#             logging.info(f"File: {filename} - Length of text: {len(text)} characters")

#             if not text:
#                 logging.warning(f"No content to process for file: {filename}")
#                 continue

#             # Convert the text into a Document object
#             documents = [Document(page_content=text)]

#             # Chunk the documents
#             chunks = chunker.chunk_documents(documents)

#             logging.info(f"File: {filename} - Number of chunks created: {len(chunks)}")

#             if not chunks:
#                 logging.warning(f"No chunks created for file: {filename}")
#                 continue

#             # Extract text from chunks
#             chunk_texts = [chunk.page_content for chunk in chunks]

#             # Generate embeddings for each chunk
#             embeddings = generate_embeddings(chunk_texts)
#             logging.info(f"File: {filename} - Generated embeddings for {len(embeddings)} chunks.")

#             if len(chunk_texts) != len(embeddings):
#                 logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
#                 continue

#              # Store chunks and embeddings in a separate file
#             save_chunks_and_embeddings(chunk_texts, embeddings, filename)   #

#             # Sanitize filename and create a document ID
#             sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '').replace('.xlsx', ''))
#             document_id = sanitized_title  # Same ID for all chunks of this document

#             # Process each chunk and create a document structure
#             for chunk_text, embedding in zip(chunk_texts, embeddings):
#                 # Validate embeddings
#                 if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
#                     logging.error(f"Embedding is not a list of floats for chunk of file: {filename}.")
#                     continue

#                 # Create a document dictionary
#                 document = {
#                     "id": document_id,
#                     "title": filename,
#                     "content": chunk_text,
#                     "contentVector": embedding
#                 }

#                 # Append processed document to the list
#                 processed_documents.append(document)

#             # Upload or index the document based on the approach
#             if config.APPROACH == 'cloud':
#                 upload_documents([document])
#                 logging.info(f"Uploaded document for {filename} to Azure Search.")
#             else:
#                 index_on_prem_faiss([embedding], [document_id])
#                 logging.info(f"Uploaded embedding for {filename} to FAISS.")

#         except Exception as e:
#             logging.error(f"Error processing file {filename}: {e}")

#     logging.info("Processing completed for all documents.")
    
#     return processed_documents  # Return the processed documents list


# # Main execution
# if __name__ == "__main__":
#     processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
#     if processed_data:
#         logging.info("Chunks and embeddings are created.")
#     else:
#         logging.error("No data was processed. Check input files or processing logic.")


# #########################################################################################

####################################################################
############ Changes #######################################################


############################################################################
import os
import json
import logging
from dotenv import load_dotenv
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import PyPDF2
import docx
import pandas as pd  # Import pandas for handling Excel files
from sentence_transformers import SentenceTransformer
from openai import AzureOpenAI
from config import Config
import re
from search_utilities import create_search_index, upload_documents

# Setup logging
logging.basicConfig(level=logging.INFO)

# Load environment variables from .env file
load_dotenv()

# Get the data folder path and chunk settings from the .env file
data_folder_path = os.getenv('DATA_FOLDER_PATH')
file_testing = "D:\\Genai_project\\Retrieval Augmented Generation\\rag_final\\RAG_task\\file_testing"
chunk_size = int(os.getenv('CHUNK_SIZE', 1500))  # Default chunk size is 1000 characters
chunk_overlap = int(os.getenv('CHUNK_OVERLAP', 300))  # Default overlap is 200 characters

config = Config()

# Initialize Azure OpenAI client if in cloud approach
if config.APPROACH == 'cloud':
    client = AzureOpenAI(azure_endpoint=os.getenv("OPENAI_API_BASE"),
                         api_key='',
                         api_version='2024-02-15-preview')
else:
    # For on-premises, initialize the sentence transformer model
    model_name = os.getenv('EMBEDDING_MODEL_NAME_ON_PREM', 'sentence-transformers/all-mpnet-base-v2')
    model = SentenceTransformer(model_name)

# Function to read PDF files
def read_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    logging.warning(f"No text found on page {page_num} of {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading PDF file {file_path}: {e}")
        return ""

# Function to read DOCX files
def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        if not text:
            logging.warning(f"No text found in DOCX file {file_path}.")
        return '\n'.join(text)
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

# Function to read TXT files
def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read().strip()
            if not text:
                logging.warning(f"No text found in TXT file {file_path}.")
            return text
    except Exception as e:
        logging.error(f"Error reading TXT file {file_path}: {e}")
        return ""



##########################################
# Function to read Excel files and chunk by rows
def read_excel(file_path):
    try:
        excel_data = pd.ExcelFile(file_path)
        text = ""
        for sheet_name in excel_data.sheet_names:
            df = excel_data.parse(sheet_name)
            # Chunk by row: Convert each row to a string and join them
            for index, row in df.iterrows():
                row_text = ' '.join(row.astype(str).values)  # Convert each cell in the row to a string
                text += f"{row_text}\n"
        return text
    except Exception as e:
        logging.error(f"Error reading Excel file {file_path}: {e}")
        return ""
    
###################################################

# Class for recursive text chunking
class RecursiveChunker:
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, docs: List[Document]) -> List[Document]:
        if not docs:
            logging.warning("No documents provided for chunking.")
            return []

        # Check if the method exists and is available
        try:
            splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                encoding_name="cl100k_base",
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
        except AttributeError as e:
            logging.error(f"Error initializing text splitter: {e}")
            return []

        # Split the documents into chunks
        chunked_docs = splitter.split_documents(docs)
        return chunked_docs

# Function to generate embeddings
def generate_embeddings(text_list, model="embedding"):
    embeddings = []
    if config.APPROACH == 'cloud':
        logging.info('Generating embeddings using cloud model...')
        for text in text_list:
            try:
                embedding = client.embeddings.create(
                    input=[text],
                    model=model  # Use the model parameter as "embedding"
                ).data[0].embedding
                embeddings.append(embedding)
            except Exception as e:
                logging.error(f"Error generating embedding for text: {e}")
    else:
        logging.info('Generating embeddings using local model...')
        try:
            embeddings = model.encode(text_list).tolist()  # Ensure embeddings are list
        except Exception as e:
            logging.error(f"Error generating embeddings using local model: {e}")
    return embeddings

####################################################
# Function to store chunks and embeddings in a separate file
def save_chunks_and_embeddings(chunks, embeddings, filename):
    try:
        output = []
        for chunk, embedding in zip(chunks, embeddings):
            chunk_data = {
                'chunk': chunk,
                'embedding': embedding
            }
            output.append(chunk_data)

        # Save to a JSON file
        output_file = os.path.join(file_testing, f"{filename}_chunks_embeddings.json")
        with open(output_file, 'w') as f:
            json.dump(output, f, indent=4)

        logging.info(f"Saved chunks and embeddings to {output_file}")
    except Exception as e:
        logging.error(f"Error saving chunks and embeddings: {e}")

###################################################

# Function to handle multiple file types, chunking, and embeddings
def process_files(data_folder_path, chunk_size, chunk_overlap):
    # Create the search index at the beginning
    create_search_index()

    chunker = RecursiveChunker(chunk_size, chunk_overlap)
    processed_documents = []  # Create a list to hold processed documents

    for filename in os.listdir(data_folder_path):
        file_path = os.path.join(data_folder_path, filename)
        logging.info(f"Processing file: {filename}")

        try:
            # Handle different file types
            if filename.endswith('.pdf'):
                text = read_pdf(file_path)
            elif filename.endswith('.docx'):
                text = read_docx(file_path)
            elif filename.endswith('.txt'):
                text = read_txt(file_path)
            elif filename.endswith('.xlsx'):  # Check if it's an Excel file
                text = read_excel(file_path)
            else:
                logging.warning(f"Unsupported file format: {filename}")
                continue

            logging.info(f"File: {filename} - Length of text: {len(text)} characters")

            if not text:
                logging.warning(f"No content to process for file: {filename}")
                continue

            # Convert the text into a Document object
            documents = [Document(page_content=text)]

            # Chunk the documents
            chunks = chunker.chunk_documents(documents)

            logging.info(f"File: {filename} - Number of chunks created: {len(chunks)}")

            if not chunks:
                logging.warning(f"No chunks created for file: {filename}")
                continue

            # Extract text from chunks
            chunk_texts = [chunk.page_content for chunk in chunks]

            # Generate embeddings for each chunk
            embeddings = generate_embeddings(chunk_texts)
            logging.info(f"File: {filename} - Generated embeddings for {len(embeddings)} chunks.")

            if len(chunk_texts) != len(embeddings):
                logging.error(f"Mismatch between number of chunks ({len(chunk_texts)}) and embeddings ({len(embeddings)})")
                continue

             # Store chunks and embeddings in a separate file
            save_chunks_and_embeddings(chunk_texts, embeddings, filename)   #

            # Sanitize filename and create a document ID
            sanitized_title = re.sub(r'[^A-Za-z0-9-_]', '_', filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '').replace('.xlsx', ''))
            document_id = sanitized_title  # Same ID for all chunks of this document

            # Process each chunk and create a document structure
            for idx, (chunk_text, embedding) in enumerate(zip(chunk_texts, embeddings)):  #
                # Validate embeddings
                if not (isinstance(embedding, list) and all(isinstance(x, float) for x in embedding)):
                    logging.error(f"Embedding is not a list of floats for chunk of file: {filename}.")
                    continue

                # Create a unique document ID for each chunk
                document_id = f"{sanitized_title}_chunk_{idx}"  #

                # Create a document dictionary
                document = {
                    "id": document_id,
                    "title": filename,
                    "content": chunk_text,
                    "contentVector": embedding
                }

                # Append processed document to the list
                processed_documents.append(document)

            # Upload or index the document based on the approach
            if config.APPROACH == 'cloud':
                upload_documents([document])
                logging.info(f"Uploaded document for {filename} to Azure Search.")
            else:
                index_on_prem_faiss([embedding], [document_id])
                logging.info(f"Uploaded embedding for {filename} to FAISS.")

        except Exception as e:
            logging.error(f"Error processing file {filename}: {e}")

    logging.info("Processing completed for all documents.")
    
    return processed_documents  # Return the processed documents list


# Main execution
if __name__ == "__main__":
    processed_data = process_files(data_folder_path, chunk_size, chunk_overlap)
    if processed_data:
        logging.info("Chunks and embeddings are created.")
    else:
        logging.error("No data was processed. Check input files or processing logic.")


#########################################################################################